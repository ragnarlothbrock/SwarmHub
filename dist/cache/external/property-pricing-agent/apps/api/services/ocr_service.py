"""OCR service for extracting text from documents.

This module provides OCR (Optical Character Recognition) capabilities
for extracting searchable text from PDF, image, and Word documents.
"""

import asyncio
import logging
import subprocess
from pathlib import Path
from typing import Optional

from core.security_utils import sanitize_for_log

logger = logging.getLogger(__name__)

# Check if OCR dependencies are available
try:
    import pytesseract

    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract not installed - OCR will be limited")

try:
    from pdf2image import convert_from_path

    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logger.warning("pdf2image not installed - PDF OCR will not work")

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed - DOCX extraction will not work")

try:
    from PIL import Image

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not installed - Image OCR will not work")


class OCRError(Exception):
    """Exception raised for OCR processing errors."""

    pass


class OCRService:
    """Service for extracting text from documents using OCR.

    Supports:
    - PDF files (via pdf2image + pytesseract)
    - Images (via pytesseract)
    - DOC/DOCX files (via python-docx)
    """

    def __init__(self, language: str = "eng"):
        """Initialize OCR service.

        Args:
            language: Tesseract language code (default: 'eng' for English)
        """
        self.language = language
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if required OCR dependencies are installed."""
        if not TESSERACT_AVAILABLE:
            logger.warning("pytesseract not available. Install with: pip install pytesseract")

        if not PDF2IMAGE_AVAILABLE:
            logger.warning("pdf2image not available. Install with: pip install pdf2image")

        if TESSERACT_AVAILABLE:
            try:
                # Check if tesseract binary is available
                result = subprocess.run(
                    ["tesseract", "--version"],
                    capture_output=True,
                    text=True,
                )
                if result.returncode == 0:
                    logger.info(
                        "Tesseract installed: %s", sanitize_for_log(result.stdout.split()[1])
                    )
                else:
                    logger.warning("Tesseract binary not found in PATH")
            except FileNotFoundError:
                logger.warning("Tesseract binary not found. Install tesseract-ocr package")

    def is_available(self) -> bool:
        """Check if OCR service is available.

        Returns:
            True if at least one OCR method is available
        """
        return TESSERACT_AVAILABLE or DOCX_AVAILABLE

    async def extract_text(
        self, file_path: str, file_type: str, max_pages: int = 50
    ) -> tuple[Optional[str], Optional[str]]:
        """Extract text from a document file.

        Args:
            file_path: Path to the document file
            file_type: MIME type of the file
            max_pages: Maximum number of pages to process (for PDFs)

        Returns:
            Tuple of (extracted_text, error_message)
        """
        path = Path(file_path)
        if not path.exists():
            return None, f"File not found: {file_path}"

        try:
            if file_type == "application/pdf":
                return await self._extract_from_pdf(file_path, max_pages)
            elif file_type in ("image/jpeg", "image/jpg", "image/png"):
                return await self._extract_from_image(file_path)
            elif file_type == "application/msword":
                return await self._extract_from_doc(file_path)
            elif (
                file_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return await self._extract_from_docx(file_path)
            else:
                return None, f"Unsupported file type for OCR: {file_type}"

        except Exception as e:
            logger.error(
                "OCR extraction failed for %s: %s", sanitize_for_log(file_path), sanitize_for_log(e)
            )
            return None, str(e)

    async def _extract_from_pdf(
        self, file_path: str, max_pages: int = 50
    ) -> tuple[Optional[str], Optional[str]]:
        """Extract text from a PDF file.

        Args:
            file_path: Path to the PDF file
            max_pages: Maximum number of pages to process

        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not PDF2IMAGE_AVAILABLE or not TESSERACT_AVAILABLE:
            return None, "PDF OCR requires pdf2image and pytesseract packages"

        try:
            # Run in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._pdf_to_text_sync, file_path, max_pages)
            return text, None
        except Exception as e:
            return None, f"PDF OCR failed: {e}"

    def _pdf_to_text_sync(self, file_path: str, max_pages: int) -> str:
        """Synchronous PDF to text conversion (runs in thread pool)."""
        # Convert PDF to images
        images = convert_from_path(
            file_path,
            dpi=200,  # Lower DPI for faster processing
            first_page=1,
            last_page=max_pages,
        )

        # Extract text from each page
        texts = []
        for i, image in enumerate(images):
            try:
                page_text = pytesseract.image_to_string(image, lang=self.language)
                if page_text.strip():
                    texts.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
            except Exception as e:
                logger.warning(
                    "Failed to OCR page %s: %s", sanitize_for_log(i + 1), sanitize_for_log(e)
                )

        return "\n\n".join(texts) if texts else ""

    async def _extract_from_image(self, file_path: str) -> tuple[Optional[str], Optional[str]]:
        """Extract text from an image file.

        Args:
            file_path: Path to the image file

        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not TESSERACT_AVAILABLE or not PIL_AVAILABLE:
            return None, "Image OCR requires pytesseract and Pillow packages"

        try:
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._image_to_text_sync, file_path)
            return text, None
        except Exception as e:
            return None, f"Image OCR failed: {e}"

    def _image_to_text_sync(self, file_path: str) -> str:
        """Synchronous image to text conversion (runs in thread pool)."""
        image = Image.open(file_path)
        return pytesseract.image_to_string(image, lang=self.language).strip()  # type: ignore[no-any-return]

    async def _extract_from_doc(self, file_path: str) -> tuple[Optional[str], Optional[str]]:
        """Extract text from a .doc file.

        Note: .doc files (old Word format) require different handling.
        For now, we return a message about limited support.

        Args:
            file_path: Path to the .doc file

        Returns:
            Tuple of (extracted_text, error_message)
        """
        # .doc files need antiword or similar tool
        # For now, return limited support message
        return None, "Legacy .doc format has limited support. Convert to .docx for better results."

    async def _extract_from_docx(self, file_path: str) -> tuple[Optional[str], Optional[str]]:
        """Extract text from a .docx file.

        Args:
            file_path: Path to the .docx file

        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not DOCX_AVAILABLE:
            return None, "DOCX extraction requires python-docx package"

        try:
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._docx_to_text_sync, file_path)
            return text, None
        except Exception as e:
            return None, f"DOCX extraction failed: {e}"

    def _docx_to_text_sync(self, file_path: str) -> str:
        """Synchronous DOCX to text conversion (runs in thread pool)."""
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return "\n\n".join(paragraphs)


# Global instance
_ocr_service: Optional[OCRService] = None


def get_ocr_service() -> OCRService:
    """Get the global OCR service instance.

    Returns:
        OCRService instance
    """
    global _ocr_service
    if _ocr_service is None:
        _ocr_service = OCRService()
    return _ocr_service
